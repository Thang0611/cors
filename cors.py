# Tạo file Word với nội dung chi tiết về CORS
from docx import Document

doc_detailed = Document()
doc_detailed.add_heading('Giải Thích Chi Tiết Về CORS (Cross-Origin Resource Sharing)', 0)

doc_detailed.add_heading('CORS là gì?', level=1)
doc_detailed.add_paragraph(
    "CORS (Cross-Origin Resource Sharing) là một cơ chế bảo mật của trình duyệt cho phép hoặc từ chối việc truy cập tài nguyên trên một trang web từ một nguồn gốc khác. "
    "Nguồn gốc được định nghĩa bằng ba thành phần: giao thức (HTTP/HTTPS), domain (tên miền), và port (cổng)."
)

doc_detailed.add_heading('Nguyên nhân cần CORS', level=1)
doc_detailed.add_paragraph(
    "Trong web, các yêu cầu giữa các nguồn gốc khác nhau thường bị chặn bởi trình duyệt để ngăn chặn các cuộc tấn công như CSRF (Cross-Site Request Forgery) và XSS (Cross-Site Scripting). "
    "Điều này có nghĩa là một trang web không thể tự do lấy dữ liệu từ một trang web khác trừ khi máy chủ của trang web đó cho phép thông qua CORS."
)

doc_detailed.add_heading('Cách CORS hoạt động', level=1)
doc_detailed.add_paragraph(
    "1. Yêu cầu đơn giản: Khi trang A gửi yêu cầu (như GET hoặc POST) đến trang B, trình duyệt sẽ gửi yêu cầu này. "
    "Nếu trang B không thiết lập CORS, trình duyệt sẽ chặn yêu cầu này.\n"
    "2. Preflight request (Yêu cầu kiểm tra trước):\n"
    "   - Nếu yêu cầu không phải là yêu cầu đơn giản, trình duyệt sẽ gửi một yêu cầu kiểm tra trước bằng phương thức OPTIONS để hỏi máy chủ.\n"
    "   - Yêu cầu này sẽ hỏi máy chủ: \"Tôi có thể gửi yêu cầu đến bạn không?\" với các thông tin như phương thức HTTP và header sẽ được sử dụng.\n"
    "3. Phản hồi từ máy chủ: Máy chủ sẽ trả lời với các header CORS như Access-Control-Allow-Origin, Access-Control-Allow-Methods, Access-Control-Allow-Headers."
)

doc_detailed.add_heading('Ví dụ minh họa', level=1)
doc_detailed.add_paragraph(
    "Giả sử bạn có một trang web trên https://example-a.com và bạn muốn truy cập dữ liệu từ https://example-b.com/api/data:\n"
    "- Yêu cầu từ trang A:\n"
    "  ```javascript\n"
    "  fetch('https://example-b.com/api/data')\n"
    "      .then(response => response.json())\n"
    "      .then(data => console.log(data));\n"
    "  ```\n"
    "- Yêu cầu kiểm tra trước (nếu cần):\n"
    "  ```http\n"
    "  OPTIONS /api/data HTTP/1.1\n"
    "  Host: example-b.com\n"
    "  Origin: https://example-a.com\n"
    "  Access-Control-Request-Method: GET\n"
    "  ```\n"
    "- Phản hồi từ trang B:\n"
    "  ```http\n"
    "  HTTP/1.1 200 OK\n"
    "  Access-Control-Allow-Origin: https://example-a.com\n"
    "  Access-Control-Allow-Methods: GET, POST\n"
    "  ```"
)

doc_detailed.add_heading('Các vấn đề liên quan đến CORS', level=1)
doc_detailed.add_paragraph(
    "- CORS không hoạt động: Nếu bạn không nhận được các header CORS đúng từ máy chủ, yêu cầu sẽ bị chặn và trình duyệt sẽ hiển thị lỗi CORS trong bảng điều khiển (console).\n"
    "- CORS với cookie: Nếu bạn cần gửi cookie trong yêu cầu, bạn cần thiết lập header Access-Control-Allow-Credentials: true từ máy chủ và credentials: 'include' trong yêu cầu từ trình duyệt.\n"
    "- Debugging CORS: Khi gặp lỗi CORS, bạn có thể sử dụng bảng điều khiển trình duyệt để xem thông tin chi tiết về yêu cầu và phản hồi."
)

doc_detailed.add_heading('Kết luận', level=1)
doc_detailed.add_paragraph(
    "CORS là một phần quan trọng trong việc xây dựng ứng dụng web hiện đại, đảm bảo rằng dữ liệu được bảo vệ và chỉ có thể truy cập bởi các nguồn được phép. "
    "Hiểu rõ cách hoạt động và cấu hình CORS sẽ giúp bạn xây dựng các ứng dụng an toàn hơn."
)

# Lưu file Word
detailed_file_path = "Giai_thich_CORS_chitiet.docx"
doc_detailed.save(detailed_file_path)

detailed_file_path
